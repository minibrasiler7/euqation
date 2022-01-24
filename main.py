#import docx
from docx import Document
from équation import Equation
from solver import Solver
### à enlever si run
#from docx.document import Document
import matplotlib.mathtext as mathtext

equations = []
equations_solving = []
solv = []
for i in range(10):
    new_Eq = Equation()
    equations_solving.append(new_Eq)
    equations.append(equations_solving[i].display_equation())
    solv.append(Solver(new_Eq.coefficient_degre(0), new_Eq.coefficient_degre(1)))
    EXPRESSION = r'$'+equations[i]+'$'
    REDUCTION = r'$'+solv[i].display_equation()+'$'
    TRIER = r'$'+solv[i].display_equation_trier()+'$'
    RESULT = r'$'+solv[i].find_solution()+'$'
    offset = mathtext.math_to_image(filename_or_obj=f"image{i}.png", s=EXPRESSION)
    red = mathtext.math_to_image(filename_or_obj=f"image{i}_reduction.png", s=REDUCTION)
    tri = mathtext.math_to_image(filename_or_obj=f"image{i}_tri.png", s=TRIER)
    res = mathtext.math_to_image(filename_or_obj=f"image{i}_result.png", s=RESULT)
    print(solv[i].find_solution())

asd = Document()
asd.add_heading("Exercices - équation du premier degré")
for i in range(10):
    asd.add_picture(f"image{i}.png")
    asd.add_paragraph()

asd.add_page_break()
asd.add_heading("Corrections - équation du premier degré ")
for i in range(10):
    asd.add_picture(f"image{i}.png")
    asd.add_picture(f"image{i}_reduction.png")
    asd.add_picture(f"image{i}_tri.png")
    asd.add_picture(f"image{i}_result.png")
    asd.add_paragraph()


asd.save("exercices.docx")


